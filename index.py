import spacy
from docx import Document
from docx.shared import Inches
from datetime import datetime

from pathlib import Path
import re
from spellchecker import SpellChecker
import os

# Load English language model
nlp = spacy.load("en_core_web_sm")
spell = SpellChecker()

def filter_text(text):
    """Filters and cleans the input text by removing misspelled words"""
    # Convert to lowercase
    text = text.lower()
    
    # Handle common contractions and abbreviations
    contractions = {
        "n't": " not",
        "'ll": " will",
        "'ve": " have",
        "'re": " are",
        "'d": " would",
        "'m": " am"
    }
    
    for contraction, expansion in contractions.items():
        text = text.replace(contraction, expansion)
    
    # Remove special characters but keep basic punctuation and apostrophes
    text = re.sub(r'[^a-z\s.,!?;:()\']', ' ', text)
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Split into words and check spelling
    words = text.split()
    valid_words = []
    
    for word in words:
        # Skip single characters (except 'a' and 'i')
        if len(word) == 1 and word not in ['a', 'i']:
            continue
            
        # Keep common abbreviations and contractions
        if word in ['mr', 'mrs', 'ms', 'dr', 'prof', 'etc', 'vs', 'eg', 'ie']:
            valid_words.append(word)
            continue
            
        # Only keep words that are spelled correctly
        if spell.correction(word) == word:
            valid_words.append(word)
    
    # Join valid words back together
    filtered_text = ' '.join(valid_words)
    
    # Clean up any remaining artifacts
    filtered_text = re.sub(r'\s+', ' ', filtered_text)
    filtered_text = filtered_text.strip()
    
    return filtered_text

def calculate_ttr_spacy(text, segment_length=30):
    """Measures the Type-Token Ratio (TTR), Moving Average Type-Token Ratio (MATTR), and Lexical Density"""
    # Filter text before processing
    filtered_text = filter_text(text)
    doc = nlp(filtered_text.lower())  # Process filtered text
    
    # Tokenize (excluding punctuation & spaces & numbers)
    words = [token.text for token in doc if token.is_alpha]  

    # Compute TTR
    unique_words = set(words)
    ttr = len(unique_words) / len(words) if words else 0
    print(f"Tokens (Total words): {len(words)}\nTypes (Unique words): {len(unique_words)}")
    
    # Compute MATTR using sliding window approach
    if len(words) < segment_length:
        mattr = ttr
    else:
        # Calculate MATTR using sliding window
        ttr_values = []
        for i in range(len(words) - segment_length + 1):
            window = words[i:i + segment_length]
            window_ttr = len(set(window)) / len(window)
            ttr_values.append(window_ttr)
        
        mattr = sum(ttr_values) / len(ttr_values)  # Average TTR over all windows

    # Compute Lexical Density
    lexical_words = [token.text for token in doc if token.is_alpha and token.pos_ in ['NOUN', 'VERB', 'ADJ', 'ADV']]
    lexical_density = len(lexical_words) / len(words) if words else 0
    print(f"Lexical words: {len(lexical_words)}")

    return ttr, mattr, len(words), len(unique_words), lexical_density

def calculate_dcr(text):
    """Calculates the ratio of dependent clauses to independent clauses in the text"""
    filtered_text = filter_text(text)
    doc = nlp(filtered_text)
    
    total_clauses = 0
    dependent_clauses = 0
    independent_clauses = 0

    print("\nAnalyzing:", text)
    print("-" * 50)

    # identify dependent clauses
    for token in doc:
        # Count dependent clauses
        if token.dep_ in ["ccomp", "advcl", "acl", "relcl", "xcomp", "pcomp"]:
            dependent_clauses += 1
            print(f"🔴 Dependent Clause: '{token.text}' (Type: {token.dep_})")
    
    # identify independent clauses by looking at sentence boundaries and conjunctions
    for sent in doc.sents:
        # Get the root of the sentence
        root = [token for token in sent if token.dep_ == "ROOT"][0]
        
        # Function to get all coordinated clauses recursively
        def get_coordinated_clauses(token):
            clauses = []
            # Check for direct coordination
            for child in token.children:
                if child.dep_ == "conj":
                    # Get the complete coordinated clause
                    clause_text = " ".join([t.text for t in child.subtree])
                    clauses.append(clause_text)
                    print(f"🟢 Coordinated Clause: '{clause_text}'")
                    # Recursively check for nested coordination
                    clauses.extend(get_coordinated_clauses(child))
            return clauses
        
        # Get all coordinated clauses
        coordinated_clauses = get_coordinated_clauses(root)
        
        # Count the main clause
        if root:
            main_clause = " ".join([t.text for t in root.subtree])
            independent_clauses += 1
            print(f"🟢 Main Clause: '{main_clause}'")
        
        # Count each coordinated clause as an independent clause
        independent_clauses += len(coordinated_clauses)
    
    total_clauses = independent_clauses + dependent_clauses

    # Avoid division by zero
    dcr = dependent_clauses / total_clauses if total_clauses > 0 else 0
    return dcr, dependent_clauses, total_clauses

def analyze_student_writings(student_folder):
    """Analyzes all writings in a student's folder"""
    results = []
    folder_path = Path(student_folder)
    
    # Get all text files in the folder
    text_files = sorted(folder_path.glob('*.txt'))
    
    for file_path in text_files:
        with open(file_path, "r", encoding="utf8") as myfile:
            text = myfile.read()
            ttr, mattr, total_words, unique_words, lexical_density = calculate_ttr_spacy(text)
            dcr, dependent_clauses, total_clauses = calculate_dcr(text)
            
            results.append({
                'file_name': file_path.name,
                'date': file_path.stem,  # Assuming filename is the date
                'ttr': ttr,
                'mattr': mattr,
                'total_words': total_words,
                'unique_words': unique_words,
                'lexical_density': lexical_density,
                'dcr': dcr,
                'dependent_clauses': dependent_clauses,
                'total_clauses': total_clauses
            })
    
    return results

def create_student_report(student_name, results):
    """Creates a Word document with analysis results for a student"""
    doc = Document()
    
    # Add title
    doc.add_heading(f'Writing Development Analysis - {student_name}', 0)
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Create table with results
    table = doc.add_table(rows=len(results) + 1, cols=10)
    table.style = 'Table Grid'
    
    # Add headers
    headers = ['Date', 'Total Words', 'Unique Words', 'TTR', 'MATTR', 'Lexical Density', 
              'DCR', 'Dependent Clauses', 'Total Clauses', 'File Name']
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    
    # Add data
    for i, result in enumerate(results, start=1):
        table.cell(i, 0).text = result['date']
        table.cell(i, 1).text = str(result['total_words'])
        table.cell(i, 2).text = str(result['unique_words'])
        table.cell(i, 3).text = f"{result['ttr']:.2f} ({result['ttr']*100:.2f}%)"
        table.cell(i, 4).text = f"{result['mattr']:.2f} ({result['mattr']*100:.2f}%)"
        table.cell(i, 5).text = f"{result['lexical_density']:.2f} ({result['lexical_density']*100:.2f}%)"
        table.cell(i, 6).text = f"{result['dcr']:.2f} ({result['dcr']*100:.2f}%)"
        table.cell(i, 7).text = str(result['dependent_clauses'])
        table.cell(i, 8).text = str(result['total_clauses'])
        table.cell(i, 9).text = result['file_name']
    
    # Save the document
    doc.save(f'{student_name}_analysis_report.docx')
    print(f"Report has been saved as '{student_name}_analysis_report.docx'")

def main():
    # Get all student folders from mydata directory
    base_path = Path("mydata")
    student_folders = [f for f in base_path.iterdir() if f.is_dir()]
    all_results = {}
    
    for folder in student_folders:
        print(f"Analyzing writings in folder: {folder.name}")
        results = analyze_student_writings(folder)
        all_results[folder.name] = results
        create_student_report(folder.name, results)
    
    # Get unique dates for all the students
    all_dates = set()
    for student_results in all_results.values():
        for result in student_results:
            all_dates.add(result['date'])
    all_dates = sorted(list(all_dates))
    
    print("Analysis complete! Check the generated reports!")

if __name__ == '__main__':
    main()
